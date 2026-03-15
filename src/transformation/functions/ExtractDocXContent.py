import mammoth
from html.parser import HTMLParser

from src.inputFiles.DocXFile import DocXFile
from src.inputFiles.TextContent import TextContent
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
import re
from docx.oxml.ns import qn
import docx 
from docx.document import Document
from docx import Document as DocumentInit
import io

def normalize_text(text: str) -> str:
    """
    Normalizes extracted text by:
    - Replacing curly quotes with straight quotes
    - Removing excessive white spaces
    - Fixing bullet points formatting
    - Ensuring correct JSON string formatting
    """
    # Normalize curly quotes
    text = text.replace("’", "'").replace("‘", "'")  # Single quotes
    text = text.replace("“", "\"").replace("”", "\"")  # Double quotes

    # Fix bullet points and spacing
    text = re.sub(r"\n\s*•\s*", "\n- ", text)  # Ensure consistent bullet format
    text = re.sub(r"\s{2,}", " ", text)  # Remove excessive spaces
    text = re.sub(r"\n+", "\n", text)  # Remove excessive newlines

    # Fix JSON-breaking line breaks inside lists
    text = re.sub(r"(\w),\n(\w)", r"\1, \2", text)  # Fix misplaced newlines in lists

    return text.strip()

def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    """
    if isinstance(parent, Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("Invalid parent type")

    for child in parent_elm.iterchildren():
        if child.tag == qn('w:p'):  # Check if it's a paragraph
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):  # Check if it's a table
            yield Table(child, parent)

def get_list_prefix( paragraph):
    """
    Detects if the paragraph is part of a bulleted or numbered list.
    Returns the corresponding bullet or number as a string.
    """
    pPr = paragraph._element.find(qn("w:pPr"))  # Get paragraph properties
    if pPr is None:
        return ""  # Not a list item

    numPr = pPr.find(qn("w:numPr"))  # Find list properties
    if numPr is None:
        return ""  # Not a list item

    ilvl = numPr.find(qn("w:ilvl"))  # Indentation level
    numId = numPr.find(qn("w:numId"))  # Numbering ID

    if numId is not None:
        level = int(ilvl.get(qn("w:val"))) if ilvl is not None else 0
        return (" " * level * 2) + "• "  # Bullet points with indentation

    return ""

def do_paragraph_thing(paragraph: Paragraph):
    """
    Extracts text from a paragraph while preserving bullets, numbering, and indentation.
    """
    text = paragraph.text.strip()
    if not text:
        return ""  # Skip empty lines

    # Detect bullets and numbering
    prefix = get_list_prefix(paragraph)

    # Handle indentation
    indent = paragraph.paragraph_format.left_indent
    tab_space = "\t" if indent and indent.pt > 0 else ""

    return f"{tab_space}{prefix}{text}\n"

def do_table_thing(table: Table, verbose=False, firstTable=False):
    """
    Extracts text from tables while keeping structured formatting.
    """
    # content = ""
    # content_set = set()

    # for row in table.rows:
    #     row_content = []
    #     for cell in row.cells:
    #         cell_text = "\n".join([p.text.strip() for p in cell.paragraphs if p.text.strip()])
    #         if cell_text and cell_text not in content_set:
    #             content_set.add(cell_text)
    #             row_content.append(cell_text)

    #     if row_content:
    #         content += " | ".join(row_content) + "\n"  # Table row formatting
    #         content += "-" * 50 + "\n"  # Horizontal separator

    # return content
    content = ""
    columns = table.columns
    content_set = set()
    if firstTable:
        for i in range(len(columns)):
            cells = columns[i].cells
            for c in cells:
                if c.paragraphs != None and c.paragraphs != []:
                    if verbose:
                        print([p.text for p in c.paragraphs])
                    text_to_add =  "\n".join([p.text  for p in c.paragraphs if p.text != ""])
                    if verbose : 
                        print(text_to_add)
                    if text_to_add not in content_set:
                        content_set.add(text_to_add)
                        content += text_to_add + "\n ------------------- \n"
    else:
        for i in range(len(table.rows)):
            cells = table.rows[i].cells
            for c in cells:
                if c.paragraphs != None and c.paragraphs != []:
                    if verbose:
                        print([p.text for p in c.paragraphs])
                    text_to_add =  "\n".join([p.text  for p in c.paragraphs if p.text != ""])
                    if verbose : 
                        print(text_to_add)
                    if text_to_add not in content_set:
                        content_set.add(text_to_add)
                        content += text_to_add + " | " 
            content += "\n ------------------- \n"

        
    return content

def get_docx_text(filename: str):
    """
    Extracts text from a DOCX file while preserving bullets, tabulations, and tables.
    """
    with open(filename, "rb") as file:
        ioBytes = io.BytesIO(file.read())
    document = DocumentInit(ioBytes)
    total_content = ""
    readTable = False

    for block_item in iter_block_items(document):
        if isinstance(block_item, Paragraph):
            total_content += do_paragraph_thing(paragraph=block_item)
        elif isinstance(block_item, Table):
            total_content += do_table_thing(table=block_item, firstTable=not readTable)
            readTable = True
    total_content = normalize_text(total_content)
    total_content = total_content.replace("\"", "")
    total_content = total_content.replace("\'", "’")
    # print(filename, "Text extracted")
    return total_content

def convert_docx_to_str_custompythondocx(docxfile: DocXFile) -> TextContent:
    operationstr = "Read docx content from file using python-docx (customFunction 1)"
    txt = get_docx_text(docxfile.filename)
    return TextContent(docxfile.id, docxfile.filename, txt, docxfile.transformationHistory + [operationstr])

class MammothHTMLParser(HTMLParser):
    """
    Parser to extract text from HTML with proper line break preservation.
    Converts block-level elements (p, li, div, h1-h6) to newline-separated text.
    """
    def __init__(self):
        super().__init__()
        self.text_parts = []
        self.current_text = ""
        self.block_elements = {'p', 'div', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'td', 'tr', 'blockquote'}
        self.previous_was_block = True
        
    def handle_starttag(self, tag, attrs):
        if tag.lower() in self.block_elements:
            # Start of block element
            if self.current_text and not self.previous_was_block:
                self.text_parts.append(self.current_text.strip())
                self.current_text = ""
            self.previous_was_block = True
    
    def handle_endtag(self, tag):
        if tag.lower() in self.block_elements:
            # End of block element - add line break
            if self.current_text.strip():
                self.text_parts.append(self.current_text.strip())
                self.current_text = ""
            self.previous_was_block = True
    
    def handle_data(self, data):
        text = data.strip()
        if text:
            if self.current_text:
                self.current_text += " " + text
            else:
                self.current_text = text
            self.previous_was_block = False
    
    def get_text(self):
        if self.current_text.strip():
            self.text_parts.append(self.current_text.strip())
        # Join with newlines and clean up excessive whitespace
        result = '\n'.join(self.text_parts)
        # Clean up excessive newlines
        result = re.sub(r'\n\s*\n+', '\n', result)
        return result.strip()

def convert_docx_to_str_mammoth(docxfile: DocXFile) -> TextContent:
    """
    Extract text from DocX using Mammoth with multi-line formatting preserved.
    
    This function:
    - Converts DocX to HTML using Mammoth
    - Parses HTML to extract text with proper line breaks
    - Preserves paragraph structure, lists, and headers
    """
    operationstr = "Read docx content from file using Mammoth"
    
    try:
        # Convert DOCX to HTML first (preserves structure better)
        result = mammoth.convert_to_html(docxfile.filename)
        html_content = result.value
        
        # Parse HTML to extract text with line breaks
        parser = MammothHTMLParser()
        parser.feed(html_content)
        txt = parser.get_text()
        
        # Final normalization
        txt = normalize_text(txt)
        
    except Exception as e:
        # Fallback to raw text extraction
        txt = mammoth.extract_raw_text(docxfile.filename).value
        txt = normalize_text(txt)
    
    return TextContent(docxfile.id, docxfile.filename, txt, docxfile.transformationHistory + [operationstr])